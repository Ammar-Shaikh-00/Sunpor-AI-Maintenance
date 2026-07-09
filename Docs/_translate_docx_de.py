from __future__ import annotations

import re
import time
import urllib.parse
from pathlib import Path

import requests
from docx import Document


SRC = Path(r"c:\Users\Ammar Shaikh\Desktop\SUNPOR_Requirements_Document.docx")
OUT = Path(r"c:\Users\Ammar Shaikh\Desktop\SUNPOR_Requirements_Document_DE.docx")
ALPHA_RE = re.compile(r"[A-Za-z]")


def should_translate(text: str) -> bool:
    t = text.strip()
    return bool(t and ALPHA_RE.search(t))


def main() -> None:
    doc = Document(str(SRC))
    cache: dict[str, str] = {}

    def tr(text: str) -> str:
        if text in cache:
            return cache[text]
        if not should_translate(text):
            cache[text] = text
            return text

        url = (
            "https://translate.googleapis.com/translate_a/single"
            "?client=gtx&sl=en&tl=de&dt=t&q="
            + urllib.parse.quote(text)
        )
        translated = text
        for attempt in range(5):
            try:
                resp = requests.get(url, timeout=12)
                resp.raise_for_status()
                data = resp.json()
                translated = "".join(part[0] for part in data[0] if part and part[0]) or text
                break
            except Exception:
                time.sleep(0.4 * (attempt + 1))
        cache[text] = translated
        return translated

    # Paragraph-level translation keeps heading/list/table formatting intact.
    for i, paragraph in enumerate(doc.paragraphs, start=1):
        text = paragraph.text
        if should_translate(text):
            paragraph.text = tr(text)
        if i % 60 == 0:
            print(f"paragraphs: {i}/{len(doc.paragraphs)}")

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    if should_translate(text):
                        paragraph.text = tr(text)

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            text = paragraph.text
            if should_translate(text):
                paragraph.text = tr(text)
        for paragraph in section.footer.paragraphs:
            text = paragraph.text
            if should_translate(text):
                paragraph.text = tr(text)

    doc.save(str(OUT))
    print(f"wrote: {OUT}")
    print(f"translated unique strings: {len(cache)}")


if __name__ == "__main__":
    main()
