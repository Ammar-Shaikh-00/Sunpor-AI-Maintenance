"""Build SUNPOR Requirements Document Word files (EN + DE) from markdown.

Usage:
    python Docs/_build_reqdoc.py

Produces:
    Docs/SUNPOR_Requirements_Document.docx
    Docs/SUNPOR_Requirements_Document_DE.docx

Minimal Markdown-to-Word converter tuned for this document
(headings, paragraphs, bullets with **bold** / *italic* /  `code`,
horizontal rules, and GitHub-style tables).
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_LINE_SPACING


HERE = Path(__file__).resolve().parent

SOURCES = [
    (HERE / "SUNPOR_Requirements_Document.md",
     HERE / "SUNPOR_Requirements_Document.docx"),
    (HERE / "SUNPOR_Requirements_Document_DE.md",
     HERE / "SUNPOR_Requirements_Document_DE.docx"),
]


INLINE_TOKEN = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)"
)


def add_inline_runs(paragraph, text: str) -> None:
    """Add runs to a paragraph, handling **bold**, *italic*, and `code`."""
    parts = [p for p in INLINE_TOKEN.split(text) if p != ""]
    if not parts:
        paragraph.add_run(text)
        return
    for part in parts:
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
        else:
            paragraph.add_run(part)


def flush_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    header = rows[0]
    body = rows[1:]
    table = doc.add_table(rows=1 + len(body), cols=len(header))
    table.style = "Light Grid Accent 1"
    for j, cell_text in enumerate(header):
        cell = table.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        add_inline_runs(p, cell_text.strip())
        for run in p.runs:
            run.bold = True
    for i, row in enumerate(body, start=1):
        for j, cell_text in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            add_inline_runs(p, cell_text.strip())
    doc.add_paragraph()


def is_table_separator(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("|"):
        return False
    inner = stripped.strip("|").strip()
    parts = [p.strip() for p in inner.split("|")]
    return all(re.fullmatch(r":?-{3,}:?", p) for p in parts) and len(parts) > 0


def parse_table_row(line: str) -> list[str]:
    stripped = line.strip().strip("|")
    return [c.strip() for c in stripped.split("|")]


def convert(md_path: Path, docx_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    # base font
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    i = 0
    n = len(lines)
    table_buffer: list[list[str]] = []
    in_table = False

    def close_table():
        nonlocal in_table, table_buffer
        if table_buffer:
            flush_table(doc, table_buffer)
        table_buffer = []
        in_table = False

    while i < n:
        raw = lines[i]
        line = raw.rstrip()

        # Fenced code block: preserve whitespace and render as monospace paragraph
        if line.strip().startswith("```"):
            i += 1
            code_lines: list[str] = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            if i < n:
                i += 1  # skip closing ```
            for code_line in code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                run = p.add_run(code_line if code_line else " ")
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            doc.add_paragraph()
            continue

        # Table detection: header row + separator row
        if (not in_table and line.strip().startswith("|")
                and i + 1 < n and is_table_separator(lines[i + 1])):
            in_table = True
            table_buffer = [parse_table_row(line)]
            i += 2  # skip header + separator
            while i < n and lines[i].strip().startswith("|"):
                table_buffer.append(parse_table_row(lines[i]))
                i += 1
            close_table()
            continue

        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if line.strip() == "---":
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2).strip()
            h = doc.add_heading(level=min(level, 4))
            add_inline_runs(h, heading_text)
            i += 1
            continue

        # Bullet lists (support "-", "*", and nested "  -" / "  1." with indent)
        m_bullet = re.match(r"^(\s*)([-*])\s+(.*)$", line)
        m_num = re.match(r"^(\s*)(\d+)\.\s+(.*)$", line)
        if m_bullet or m_num:
            m = m_bullet or m_num
            indent = len(m.group(1))
            content = m.group(3)
            style_name = "List Bullet" if m_bullet else "List Number"
            # Word doesn't do nested levels without a lot of XML. Use indent visually.
            para = doc.add_paragraph(style=style_name)
            if indent >= 2:
                para.paragraph_format.left_indent = Pt(18 + (indent // 2) * 12)
            add_inline_runs(para, content)
            i += 1
            continue

        # Regular paragraph (may span multiple lines until blank)
        buf = [line]
        i += 1
        while i < n and lines[i].strip() and not re.match(
                r"^(#{1,6}\s|[-*]\s|\d+\.\s|\|)", lines[i].lstrip()):
            buf.append(lines[i].rstrip())
            i += 1
        para_text = " ".join(s.strip() for s in buf)
        p = doc.add_paragraph()
        add_inline_runs(p, para_text)

    close_table()
    doc.save(docx_path)
    print(f"wrote {docx_path.name}")


def main() -> int:
    for md, docx in SOURCES:
        if not md.exists():
            print(f"missing source: {md}", file=sys.stderr)
            return 1
        convert(md, docx)
    return 0


if __name__ == "__main__":
    sys.exit(main())
