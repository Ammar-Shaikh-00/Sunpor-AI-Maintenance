from __future__ import annotations

import re
import urllib.parse
from pathlib import Path

import requests
from docx import Document


SRC = Path(r"c:\Users\Ammar Shaikh\Desktop\SUNPOR_Requirements_Document.docx")
OUT = Path(r"c:\Users\Ammar Shaikh\Desktop\SUNPOR_Requirements_Document_DE_v2.docx")

ALPHA_RE = re.compile(r"[A-Za-z]")


def should_translate(text: str) -> bool:
    t = text.strip()
    return bool(t and ALPHA_RE.search(t))


def translate_en_to_de(text: str, cache: dict[str, str]) -> str:
    if text in cache:
        return cache[text]
    if not should_translate(text):
        cache[text] = text
        return text

    url = (
        "https://translate.googleapis.com/translate_a/single"
        "?client=gtx&sl=en&tl=de&dt=t&q="
        + urllib.parse.quote(text)
    )
    try:
        resp = requests.get(url, timeout=12)
        resp.raise_for_status()
        data = resp.json()
        translated = "".join(part[0] for part in data[0] if part and part[0]) or text
    except Exception:
        translated = text
    cache[text] = translated
    return translated


def set_paragraph_text_keep_layout(paragraph, new_text: str) -> None:
    # Keep paragraph/list/table style and alignment unchanged.
    # Replace run text content only.
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def main() -> None:
    doc = Document(str(SRC))
    cache: dict[str, str] = {}

    total_paragraphs = len(doc.paragraphs)
    for i, paragraph in enumerate(doc.paragraphs, start=1):
        src_text = paragraph.text
        if should_translate(src_text):
            de_text = translate_en_to_de(src_text, cache)
            set_paragraph_text_keep_layout(paragraph, de_text)
        if i % 50 == 0:
            print(f"paragraphs translated: {i}/{total_paragraphs}")

    table_paragraphs = 0
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    src_text = paragraph.text
                    if should_translate(src_text):
                        de_text = translate_en_to_de(src_text, cache)
                        set_paragraph_text_keep_layout(paragraph, de_text)
                    table_paragraphs += 1
                    if table_paragraphs % 50 == 0:
                        print(f"table paragraphs translated: {table_paragraphs}")

    for section in doc.sections:
        for paragraph in section.header.paragraphs:
            src_text = paragraph.text
            if should_translate(src_text):
                de_text = translate_en_to_de(src_text, cache)
                set_paragraph_text_keep_layout(paragraph, de_text)
        for paragraph in section.footer.paragraphs:
            src_text = paragraph.text
            if should_translate(src_text):
                de_text = translate_en_to_de(src_text, cache)
                set_paragraph_text_keep_layout(paragraph, de_text)

    doc.save(str(OUT))
    print(f"wrote: {OUT}")
    print(f"unique strings translated: {len(cache)}")


if __name__ == "__main__":
    main()
